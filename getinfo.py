from googleapiclient.errors import HttpError
from googleapiclient.discovery import build
import google.auth
from gsheetsdb import connect
from google.oauth2 import service_account
import streamlit as st
import pandas as pd
import plotly.figure_factory as ff
import plotly.graph_objects as go
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import plotly.io as pio
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from kaleido.scopes.plotly import PlotlyScope
from PIL import Image 


# Google connection
credentials = service_account.Credentials.from_service_account_info(
    st.secrets["gcp_service_account"],
    scopes=[
        "https://www.googleapis.com/auth/spreadsheets",
    ],
)
conn = connect(credentials=credentials)
sheet_url = st.secrets["private_gsheets_url"]
sheet_id = sheet_url.split("/")[5]

# Perform SQL query on the Google Sheet.
# Uses st.cache_data to only rerun when the query changes or after 10 min.


# @st.cache_data(ttl=600)
def run_query(query):
    rows = conn.execute(query, headers=1)
    rows = rows.fetchall()
    return rows


st.title("Resultados Test 5 Disfunciones")

rows = run_query(f'SELECT * FROM "{sheet_url}"')
df = pd.DataFrame(rows)
unique_empresas = df['empresa'].unique()
# Converting the unique values into a list
unique_empresas_list = list(unique_empresas)
with st.sidebar:
    empresa = st.multiselect(
        "Elige las empresas para las que deseas generar el reporte", unique_empresas_list
    )
    
    
# Join the elements in the list with a comma and a space
empresa_str = ', '.join(empresa)
st.subheader('Empresas seleccionadas para generar reporte: ' + empresa_str)
# generate a filtered data frame with the companies selected
filtered_df = df[df['empresa'].isin(empresa)]
filtered_df_numeric = filtered_df.drop(['empresa', 'q17', 'q18', 'q20', 'q21'], axis=1)
average_values = filtered_df_numeric.mean()
confianza = average_values['q4'] + average_values['q6'] + average_values['q12']
conflicto = average_values['q1'] + average_values['q7'] + average_values['q10']
compromiso = average_values['q3'] + \
    average_values['q8'] + average_values['q13']
responsabilidad = average_values['q2'] + \
    average_values['q11'] + average_values['q14']
enfoque_resultados = average_values['q5'] + \
    average_values['q9'] + average_values['q15']

st.dataframe(filtered_df)

col1, col2, col3, col4, col5 = st.columns(5)
col1.metric("Confianza", confianza)
col2.metric("Conflicto", conflicto)
col3.metric("Compromiso", compromiso)
col4.metric("Responsabilidad", responsabilidad)
col5.metric("Enfoque en resultados", enfoque_resultados)

st.subheader("Resultados")
# Bar plot
x_values = [enfoque_resultados, responsabilidad,
            compromiso, conflicto, confianza]
# Format the text to display only one decimal place
formatted_text = [format(value, '.1f') for value in x_values]

# Define a function to determine the font color based on the value


def font_color(value):
    if value < 6:
        return 'red'
    elif value < 8:
        return 'orange'
    else:
        return 'green'


# Create a list of font colors based on the x_values
font_colors = [font_color(value) for value in x_values]

fig_resultados = go.Figure(go.Bar(
    x=x_values,
    y=['Desatención a<br>Resultados', 'Evasión de<br>Responsabilidad',
            'Falta de<br>Compromiso', 'Miedo al<br>Conflicto', 'Ausencia de<br>Confianza'],
    orientation='h',
    marker=dict(color='rgb(245,91,35)'),
    text=formatted_text,
    textfont=dict(size=17, color=font_colors, family='Montserrat'),
    textposition='outside'))
fig_resultados.update_xaxes(
    title_text="", range=[0, 9], tickvals=list(range(10)))
fig_resultados.update_layout(
    margin=dict(
        l=150,  # Adjust this value to fit your labels
        r=50,
        b=50,
        t=50,
        pad=4
    ),
)
# Plot!
st.plotly_chart(fig_resultados, use_container_width=True)

# World Cloud
st.subheader("Palabras que describen la interacción del equipo")

# Function to generate a word cloud


def generate_wordcloud(text):
    wordcloud = WordCloud(
        width=800, height=800, background_color='white', min_font_size=10).generate(text)
    return wordcloud


words = ' '.join(filtered_df['q17'])
wordcloud = generate_wordcloud(words)
fig_w, ax = plt.subplots(figsize=(10, 10))
ax.imshow(wordcloud, interpolation='bilinear')
ax.axis("off")
st.pyplot(fig_w)



# Bar plot
st.subheader('Tendencia del equipo de abordar el conflicto')

# Define the categories
categories = ["Frecuente - Irrelevante", "Poco frecuente - Irrelevante",
              "Frecuente - Relevante", "Poco Frecuente - Relevantes"]

# Calculate value counts for the 'category' column and reindex to include all categories
value_counts = filtered_df['q18'].value_counts().reindex(
    categories, fill_value=0)

# Create a bar plot with the content of the 'category' column as x-axis labels and value counts as y-axis values
fig_tendencia = go.Figure(go.Bar(
    x=value_counts.index,
    y=value_counts.values,
    marker=dict(color='rgb(245,91,35)'),
    # text=value_counts.values,
    # textposition='auto'
))
fig_tendencia.update_layout(
    margin=dict(
        l=50,  # Adjust this value to fit your labels
        r=90,
        b=200,
        t=50,
        pad=3
    ),
)
# Set x-axis title
fig_tendencia.update_xaxes(title_text="", tickangle=45, tickfont=dict(size=17))

# Set y-axis title
fig_tendencia.update_yaxes(title_text="", tickfont=dict(size=17))

# Plot!
st.plotly_chart(fig_tendencia, use_container_width=True)

# Circle graphs
st.header('Ausencia de Confianza')

# Pregunta 4, 6, 12

# Value count for the 'q1' column
vc_q4 = filtered_df['q4'].value_counts().sort_index()
vc_q6 = filtered_df['q6'].value_counts().sort_index()
vc_q12 = filtered_df['q12'].value_counts().sort_index()

# Map the values to their corresponding labels
labels = {1: 'Rara vez', 2: 'Algunas veces', 3: 'regularmente'}
vc_q4.index = vc_q4.index.map(labels)
vc_q6.index = vc_q6.index.map(labels)
vc_q12.index = vc_q12.index.map(labels)

# Define a function to generate a color list based on unique values


def generate_colors(value_counts):
    base_colors = {1: 'red', 2: 'yellow', 3: 'green'}
    # Map the index back to numeric values
    labels_to_numbers = {v: k for k, v in labels.items()}
    value_counts.index = value_counts.index.map(labels_to_numbers)
    return [base_colors[val] for val in value_counts.index]


colors_q4 = generate_colors(vc_q4)
colors_q6 = generate_colors(vc_q6)
colors_q12 = generate_colors(vc_q12)

# Create a pie chart using Plotly with custom colors
fig_4 = go.Figure(
    data=[go.Pie(labels=vc_q4.index, values=vc_q4, hole=0.6, marker=dict(colors=colors_q4),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_4.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_4)
percentage_4 = ""
total = sum(vc_q4.values)
for label, value in vc_q4.items():
    percentage_4 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)

st.write("Los miembros del equipo se disculpan rápida y genuinamente cuando dicen o hacen algo inapropiado o que pueda lastimar al equipo")
st.divider()

# Pregunta 6
# Create a pie chart using Plotly with custom colors
fig_6 = go.Figure(
    data=[go.Pie(labels=vc_q6.index, values=vc_q6, hole=0.6, marker=dict(colors=colors_q6),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_6.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_6)
percentage_6 = ""
total = sum(vc_q6.values)
for label, value in vc_q6.items():
    percentage_6 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo admiten abiertamente sus debilidades y errores")
st.divider()

# Pregunta 12
# Create a pie chart using Plotly with custom colors
fig_12 = go.Figure(
    data=[go.Pie(labels=vc_q12.index, values=vc_q12, hole=0.6, marker=dict(colors=colors_q12),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_12.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_12)
percentage_12 = ""
total = sum(vc_q12.values)
for label, value in vc_q12.items():
    percentage_12 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo conocen detalles de la vida personal de sus compañeros y platican abiertamente sobre ello")
st.divider()


# Circle graphs
st.header('Miedo al conflicto')

# Pregunta 1, 7, 10

# Value count for the 'q1' column
vc_q1 = filtered_df['q1'].value_counts().sort_index()
vc_q7 = filtered_df['q7'].value_counts().sort_index()
vc_q10 = filtered_df['q10'].value_counts().sort_index()

# Map the values to their corresponding labels
labels = {1: 'Rara vez', 2: 'Algunas veces', 3: 'regularmente'}
vc_q1.index = vc_q1.index.map(labels)
vc_q7.index = vc_q7.index.map(labels)
vc_q10.index = vc_q10.index.map(labels)

colors_q1 = generate_colors(vc_q1)
colors_q7 = generate_colors(vc_q7)
colors_q10 = generate_colors(vc_q10)

# Create a pie chart using Plotly with custom colors
fig_1 = go.Figure(
    data=[go.Pie(labels=vc_q1.index, values=vc_q1, hole=0.6, marker=dict(colors=colors_q1),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_1.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_1)
percentage_1 = ""
total = sum(vc_q1.values)
for label, value in vc_q1.items():
    percentage_1 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo son abiertos y apasionados al discutir temas laborales")
st.divider()

# Pregunta 7
# Create a pie chart using Plotly with custom colors
fig_7 = go.Figure(
    data=[go.Pie(labels=vc_q7.index, values=vc_q7, hole=0.6, marker=dict(colors=colors_q7),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_7.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_7)
percentage_7 = ""
total = sum(vc_q7.values)
for label, value in vc_q7.items():
    percentage_7 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Las juntas en la empresa son interesantes y no aburridas")
st.divider()

# Pregunta 10
# Create a pie chart using Plotly with custom colors
fig_10 = go.Figure(
    data=[go.Pie(labels=vc_q10.index, values=vc_q10, hole=0.6, marker=dict(colors=colors_q10),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_10.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_10)
percentage_10 = ""
total = sum(vc_q10.values)
for label, value in vc_q10.items():
    percentage_10 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Durante las juntas, los temas más importantes y complicados se ponen en la mesa para ser discutidos y resueltos")
st.divider()


# Circle graphs
st.header('Falta de Compromiso')

# Pregunta 3, 8, 13

# Value count for the 'q1' column
vc_q3 = filtered_df['q3'].value_counts().sort_index()
vc_q8 = filtered_df['q8'].value_counts().sort_index()
vc_q13 = filtered_df['q13'].value_counts().sort_index()

# Map the values to their corresponding labels
labels = {1: 'Rara vez', 2: 'Algunas veces', 3: 'regularmente'}
vc_q3.index = vc_q3.index.map(labels)
vc_q8.index = vc_q8.index.map(labels)
vc_q13.index = vc_q13.index.map(labels)

colors_q3 = generate_colors(vc_q3)
colors_q8 = generate_colors(vc_q8)
colors_q13 = generate_colors(vc_q13)

# Create a pie chart using Plotly with custom colors
fig_3 = go.Figure(
    data=[go.Pie(labels=vc_q3.index, values=vc_q3, hole=0.6, marker=dict(colors=colors_q3),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_3.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_3)
percentage_3 = ""
total = sum(vc_q3.values)
for label, value in vc_q3.items():
    percentage_3 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo saben en que están trabajando sus compañeros y cómo contribuyen en el logro de las metas")
st.divider()

# Pregunta 8
# Create a pie chart using Plotly with custom colors
fig_8 = go.Figure(
    data=[go.Pie(labels=vc_q8.index, values=vc_q8, hole=0.6, marker=dict(colors=colors_q8),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_8.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_8)
percentage_8 = ""
total = sum(vc_q8.values)
for label, value in vc_q8.items():
    percentage_8 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo salen de las juntas seguros de que sus compañeros están comprometidos con los acuerdos, aún cuando en la discusión estaban en desacuerdo")
st.divider()

# Pregunta 13
# Create a pie chart using Plotly with custom colors
fig_13 = go.Figure(
    data=[go.Pie(labels=vc_q13.index, values=vc_q13, hole=0.6, marker=dict(colors=colors_q13),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_13.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_13)
percentage_13 = ""
total = sum(vc_q13.values)
for label, value in vc_q13.items():
    percentage_13 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo terminan las discusiones con acuerdos claros y específicos y con planes de acción")
st.divider()

# Circle graphs
st.header('Evasión de Responsabilidades')

# Pregunta 2, 11, 14
vc_q2 = filtered_df['q2'].value_counts().sort_index()
vc_q11 = filtered_df['q11'].value_counts().sort_index()
vc_q14 = filtered_df['q14'].value_counts().sort_index()

# Map the values to their corresponding labels
labels = {1: 'Rara vez', 2: 'Algunas veces', 3: 'regularmente'}
vc_q2.index = vc_q2.index.map(labels)
vc_q11.index = vc_q11.index.map(labels)
vc_q14.index = vc_q14.index.map(labels)

colors_q2 = generate_colors(vc_q2)
colors_q11 = generate_colors(vc_q11)
colors_q14 = generate_colors(vc_q14)

# Create a pie chart using Plotly with custom colors
fig_2 = go.Figure(
    data=[go.Pie(labels=vc_q2.index, values=vc_q2, hole=0.6, marker=dict(colors=colors_q2),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_2.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_2)
percentage_2 = ""
total = sum(vc_q2.values)
for label, value in vc_q2.items():
    percentage_2 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo expresan las deficiencias y bajas de productividad de sus compañeros")
st.divider()

# Pregunta 11
# Create a pie chart using Plotly with custom colors
fig_11 = go.Figure(
    data=[go.Pie(labels=vc_q11.index, values=vc_q11, hole=0.6, marker=dict(colors=colors_q11),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_11.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_11)
percentage_11 = ""
total = sum(vc_q11.values)
for label, value in vc_q11.items():
    percentage_11 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo se preocupan por no fallarle a sus compañeros")
st.divider()

# Pregunta 14
# Create a pie chart using Plotly with custom colors
fig_14 = go.Figure(
    data=[go.Pie(labels=vc_q14.index, values=vc_q14, hole=0.6, marker=dict(colors=colors_q14),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_14.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_14)
percentage_14 = ""
total = sum(vc_q14.values)
for label, value in vc_q14.items():
    percentage_14 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo desafían los planes y formas de sus compañeros")
st.divider()

# Circle graphs
st.header('Desatención a Resultados')

# Pregunta 5, 9, 15
vc_q5 = filtered_df['q5'].value_counts().sort_index()
vc_q9 = filtered_df['q9'].value_counts().sort_index()
vc_q15 = filtered_df['q15'].value_counts().sort_index()

# Map the values to their corresponding labels
labels = {1: 'Rara vez', 2: 'Algunas veces', 3: 'regularmente'}
vc_q5.index = vc_q5.index.map(labels)
vc_q9.index = vc_q9.index.map(labels)
vc_q15.index = vc_q15.index.map(labels)

colors_q5 = generate_colors(vc_q5)
colors_q9 = generate_colors(vc_q9)
colors_q15 = generate_colors(vc_q15)

# Create a pie chart using Plotly with custom colors
fig_5 = go.Figure(
    data=[go.Pie(labels=vc_q5.index, values=vc_q5, hole=0.6, marker=dict(colors=colors_q5),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_5.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_5)
percentage_5 = ""
total = sum(vc_q5.values)
for label, value in vc_q5.items():
    percentage_5 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo están dispuestos a hacer sacrificios (presupuesto, territorio, plantilla) en sus departamentos por el bienestar del equipo")
st.divider()

# Pregunta 9
# Create a pie chart using Plotly with custom colors
fig_9 = go.Figure(
    data=[go.Pie(labels=vc_q9.index, values=vc_q9, hole=0.6, marker=dict(colors=colors_q9),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_9.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_9)
percentage_9 = ""
total = sum(vc_q9.values)
for label, value in vc_q9.items():
    percentage_9 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("El equipo se desmotiva cuando no se consiguen las metas grupales")
st.divider()

# Pregunta 15
# Create a pie chart using Plotly with custom colors
fig_15 = go.Figure(
    data=[go.Pie(labels=vc_q15.index, values=vc_q15, hole=0.6, marker=dict(colors=colors_q15),
                 textposition=None)])

# Adjust the layout, hide the legend, and set the chart size
fig_15.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=False,
    width=450,  # Adjust the width as needed
    height=300  # Adjust the height as needed
)

st.plotly_chart(fig_15)
percentage_15 = ""
total = sum(vc_q15.values)
for label, value in vc_q15.items():
    percentage_15 += f"{(value / total) * 100:.1f}% {labels[label]}\n"
    st.write(
        f"{(value / total) * 100:.1f}% {labels[label]}", unsafe_allow_html=True)
st.write("Los miembros del equipo son más rápidos para dar crédito sobre los logros de los demás que sobre los propios")
st.divider()

st.header('eNPS')

r_enps = len(filtered_df['q19'])
# Promoters are those with a score of 9 or 10
promoters = len(filtered_df[filtered_df['q19'] >= 9])
# Passives are those with a score of 7 or 8
passives = len(filtered_df[(filtered_df['q19'] >= 7) & (filtered_df['q19'] <= 8)])
# Detractors are those with a score of 6 or lower
detractors = len(filtered_df[filtered_df['q19'] <= 6])
# Calculate the percentage of promoters and detractors
percentage_promoters = (promoters / r_enps) * 100
percentage_detractors = (detractors / r_enps) * 100
percentage_passives = (passives / r_enps) * 100
# Calculate the eNPS
enps = percentage_promoters - percentage_detractors

promoters_str = str(promoters)
passives_str = str(passives)
detractors_str = str(detractors)
r_enps_str = str(r_enps)
enps_str = str(enps)

st.write('Total de respuestas eNPS: ', r_enps_str)
st.write('Total de promotores: ', promoters_str)
st.write('Total de pasivos: ', passives_str)
st.write('Total detractores: ', detractors_str)
st.write('eNPS: ', enps_str)

labels = ['Promotores', 'Pasivos', 'Detractores']
values = [percentage_promoters, percentage_passives, percentage_detractors]
colors = ['green', 'yellow', 'red']

fig_enps = go.Figure(data=[go.Pie(labels=labels, values=values, hole=.6, marker_colors=colors, textfont=dict(size=18))])

fig_enps.update_layout(
    margin=dict(l=0, r=0, t=0, b=0),
    showlegend=True,
    annotations=[dict(text='eNPS', x=0.5, y=0.5, font_size=20, showarrow=False)],
    legend=dict(x=0.1, y=1.1, font=dict(size=16))
)
st.plotly_chart(fig_enps)

st.subheader('Principales Razones')
# Extracting text from 'q20' column and formatting as an unordered list
razones_list = filtered_df['q20'].tolist()
bullet_razones = "\n".join(f"• {item}" for item in razones_list)
st.write(bullet_razones)

st.subheader('Iniciativas')
# Extracting text from 'q21' column and formatting as an unordered list
comportamiento_list = filtered_df['q21'].tolist()
bullet_comportamiento = "\n".join(f"• {item}" for item in comportamiento_list)
st.write(bullet_comportamiento)

empresa_str = ' '.join(empresa)

fig_resultados.write_image("fig_resultados.png", width=700, height=420)
fig_4.write_image("fig_4.png", width=700, height=420)
fig_6.write_image("fig_6.png", width=700, height=420)
fig_12.write_image("fig_12.png", width=700, height=420)
fig_1.write_image("fig_1.png", width=700, height=420)
fig_7.write_image("fig_7.png", width=700, height=420)
fig_10.write_image("fig_10.png", width=700, height=420)
fig_3.write_image("fig_3.png", width=700, height=420)
fig_8.write_image("fig_8.png", width=700, height=420)
fig_13.write_image("fig_13.png", width=700, height=420)
fig_2.write_image("fig_2.png", width=700, height=420)
fig_11.write_image("fig_11.png", width=700, height=420)
fig_14.write_image("fig_14.png", width=700, height=420)
fig_13.write_image("fig_13.png", width=700, height=420)
fig_5.write_image("fig_5.png", width=700, height=420)
fig_9.write_image("fig_9.png", width=700, height=420)
fig_15.write_image("fig_15.png", width=700, height=420)
wordcloud.to_file("wordcloud.png")
fig_enps.write_image("fig_enps.png", width=700, height=420)
fig_tendencia.write_image("fig_tendencia.png", width=700, height=420)



def replace_placeholder_with_image(paragraph, placeholder, image_path, width_inches=3):
    if placeholder in paragraph.text:
        # Split the paragraph text at the placeholder
        before, placeholder, after = paragraph.text.partition(placeholder)

        # Clear the paragraph and add new Run objects
        paragraph.clear()

        # Add the text before the placeholder
        if before:
            paragraph.add_run(before)

        # Open the image with Pillow and get its dimensions
        image = Image.open(image_path)
        width, height = image.size

        # Calculate the proportional height based on the desired width
        aspect_ratio = float(height) / float(width)
        height_inches = width_inches * aspect_ratio

        # Add the image with the desired width and proportional height
        paragraph.add_run().add_picture(image_path, width=Inches(width_inches), height=Inches(height_inches))

        # Add the text after the placeholder
        if after:
            paragraph.add_run(after)

def replace_placeholder_with_image_in_table(table, placeholder, image_path, width_inches=3):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholder_with_image(paragraph, placeholder, image_path, width_inches)


def replace_placeholder_with_text(paragraph, placeholder, text):
    if placeholder in paragraph.text:
        for run in paragraph.runs:
            run.text = run.text.replace(placeholder, text)

def replace_placeholder_with_text_in_table(table, placeholder, text):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholder_with_text(paragraph, placeholder, text)


# Open the existing Word document
doc = Document('Disfunciones de un equipo GTC.docx')

# Replace placeholders with images
for paragraph in doc.paragraphs:
    replace_placeholder_with_image(paragraph, '{fig_resultados}', 'fig_resultados.png', width_inches=6.5)

for table in doc.tables:
    replace_placeholder_with_image_in_table(table, '{fig_4}', 'fig_4.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_6}', 'fig_6.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_12}', 'fig_12.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_1}', 'fig_1.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_7}', 'fig_7.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_10}', 'fig_10.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_3}', 'fig_3.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_8}', 'fig_8.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_13}', 'fig_13.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_2}', 'fig_2.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_11}', 'fig_11.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_14}', 'fig_14.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_5}', 'fig_5.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_9}', 'fig_9.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{fig_15}', 'fig_15.png', width_inches=3)
    replace_placeholder_with_image_in_table(table, '{wordcloud}', 'wordcloud.png', width_inches=5)
    replace_placeholder_with_text_in_table(table, '{percentage_4}', percentage_4)
    replace_placeholder_with_text_in_table(table, '{percentage_6}', percentage_6)
    replace_placeholder_with_text_in_table(table, '{percentage_12}', percentage_12)
    replace_placeholder_with_text_in_table(table, '{percentage_1}', percentage_1)
    replace_placeholder_with_text_in_table(table, '{percentage_7}', percentage_7)
    replace_placeholder_with_text_in_table(table, '{percentage_10}', percentage_10)
    replace_placeholder_with_text_in_table(table, '{percentage_3}', percentage_3)
    replace_placeholder_with_text_in_table(table, '{percentage_8}', percentage_8)
    replace_placeholder_with_text_in_table(table, '{percentage_13}', percentage_13)
    replace_placeholder_with_text_in_table(table, '{percentage_2}', percentage_2)
    replace_placeholder_with_text_in_table(table, '{percentage_11}', percentage_11)
    replace_placeholder_with_text_in_table(table, '{percentage_14}', percentage_14)
    replace_placeholder_with_text_in_table(table, '{percentage_5}', percentage_5)
    replace_placeholder_with_text_in_table(table, '{percentage_9}', percentage_9)
    replace_placeholder_with_text_in_table(table, '{percentage_15}', percentage_15)
    replace_placeholder_with_text_in_table(table, '{promoters}', promoters_str)
    replace_placeholder_with_text_in_table(table, '{passives}', passives_str)
    replace_placeholder_with_text_in_table(table, '{detractors}', detractors_str)
    replace_placeholder_with_text_in_table(table, '{enps}', enps_str)
    replace_placeholder_with_image_in_table(table, '{fig_enps}', 'fig_enps.png', width_inches=4)
    replace_placeholder_with_image_in_table(table, '{fig_tendencia}', 'fig_tendencia.png', width_inches=6)

# Replace placeholders with text
for paragraph in doc.paragraphs:
    replace_placeholder_with_text(paragraph, '{empresa}', empresa_str)
    if '{razones}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{razones}', bullet_razones)
    if '{comportamiento}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{comportamiento}', bullet_comportamiento)      

# Save the Word document with the new content
doc.save('updated_word_document.docx')

with st.sidebar:
    
    # # Create a button, but only if the code has run successfully
    # Read the .docx file as binary
    with open('updated_word_document.docx', 'rb') as file:
        bytes_data = file.read()
    
    st.download_button(
        label="Download report in word",
        data=bytes_data,
        file_name="5dis_"+f"{empresa_str}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )